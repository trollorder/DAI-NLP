import docx
import pandas as pd
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
import os
from datetime import date
from dalleimggen import get_img_prompt
import assess_sustainability

def get_report(report_path,product,globalvalues={}):
    product = "apple watch"
    
    #global variables dafault and will update if any values
    numberofcomments = globalvalues["numberofcomments"]
    componentliststring = globalvalues["componentliststring"]
    promptliststring = globalvalues["promptliststring"]
    globalsentiment = globalvalues["globalsentiment"]
    accuracy = globalvalues["accuracy"]
    today = date.today()
    
    # create a new document
    doc = docx.Document()
    

    #checks number of table
    tablecount = 0

    #get list of recommendation
    recco = pd.read_csv(f"{report_path}/Recommendation.csv")
    recco["Sentiment Cumulative Score"] = recco[["Sentiment Cumulative Score"]].apply(lambda x : abs(x))
    recco.sort_values(by='Sentiment Cumulative Score', inplace=True)
    print(recco["Sentiment Cumulative Score"])
    get_img_prompt(report_path,product,list(recco["Recommendation"]))

    #Styling

    # set the default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = docx.shared.Pt(12)

    # add a header to the document
    header = doc.sections[0].header
    header.paragraphs[0].style = 'Heading 1'
    header.paragraphs[0].text = f"AID Report - Sentiment Analysis and Report for {product} "

    #add hero img which is the first image
    listofimg = os.listdir(f"{report_path}/img")
    print(listofimg)
    #get relevant file
    relevantnames = []
    for imagename in listofimg:
        if product in imagename:
            relevantnames.append(imagename)
            
    print(relevantnames)
    try:
        doc.add_picture(f'img/{product}1.png', width=Inches(3))
    except:
        print("No header image found")
    lastp = doc.paragraphs[-1]
    # Set the alignment of the paragraph to center
    lastp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # add a table to the document
    def addtable(doc,tableofvals,tablecount,headerlist,cellwidth = []):
        rowintable = len(tableofvals)
        colintable = len(tableofvals[0])
        table = doc.add_table(rows=1,cols = len(headerlist))
        #set headers 
        heading_cells = table.rows[0].cells
        for idx,x in enumerate(headerlist):
            heading_cells[idx].text = str(x)
        table = doc.tables[tablecount]
        table.style = "Table Grid"
        
        #set values
        print(f"There is {len(tableofvals)} number of rows in this table")
        for rowidx , rowdata in enumerate(tableofvals):
            new_row = table.add_row()
            for colidx,cellvalue in enumerate(rowdata):

                new_row.cells[colidx].text = str(cellvalue)
        #set widths
        for index,widths in enumerate(cellwidth):
            if len(cellwidth) == 0:
                break
            else:
                for y in range(len(tableofvals)+1):
                    table.cell(y,index).width = Cm(widths)        
                
        tablecount += 1
        return tablecount
        
    
    
    tabletostore= [["Product Name" ,product],
                ["Number of Comments Checked" , numberofcomments],
                ["Overall Product Sentiment" , globalsentiment],
                ["Components Analysed" , componentliststring],
                ["Prompts Used" , promptliststring],
                ["Model Accuracy" , accuracy],
                ["Date Generated", today ]]

    #return variable tracks which table is referenced next
    tablecount = addtable(doc,tabletostore,tablecount,["Product Fields" , "Details"] )

    #add graph here
    doc.add_heading("Sentiment Graph Over Time", level=2)
    doc.add_picture(f'img/{product} Sentiment Analysis Overtime.png', width=Inches(6))

    #component, recommendation, score


    #dyamically inject rows
    reccotable = [[] for x in range(len(recco)) ]
    for idx,row in recco.iterrows():
        newrow= [str(row["Component Specification"]),str(row["Recommendation"]),str(row["Sentiment Cumulative Score"]) , str(row["Benchmark"])]
        reccotable[idx] += newrow
        print(newrow)
        if idx>20: #table limit
            break
    doc.add_heading("Table of Recommendations" , level = 2)
    tablecount = addtable(doc,reccotable,tablecount,["Component Specification" , "Recommendation" , "Priority" , "Current Benchmark"],[3,7,2,3])
    #sorting code below
    table = doc.tables[tablecount-1]
    # Iterate over all the rows in the table to clear empty rows
    new_table = docx.Document().add_table(rows=0, cols=len(table.columns))
    new_table.style = "Table Grid"
    for row in table.rows:
        # Check if the row is empty
        if not all(cell.text.strip() == '' for cell in row.cells):
            # If the row is not empty, copy it to the new table
            new_row = new_table.add_row()
            for i, cell in enumerate(row.cells):
                new_row.cells[i].text = cell.text
    cellwidth = [3,7,2,3]
    for x,row in enumerate(new_table.rows):
        for y,col in enumerate(new_table.columns):
            new_table.cell(x,y).width = Cm(cellwidth[y])     
    # Replace the original table with the new table
    table._element.getparent().replace(table._element, new_table._element)
    
    nameusedlist = list(pd.read_csv(f"{report_path}/{product} dalle prompts used.csv")["imgname"])
    promptusedlist = list(pd.read_csv(f"{report_path}/{product} dalle prompts used.csv")["Prompts used for dall e"])
    doc.add_heading("Table of Dall E generated images and associated prompts" , level = 2)
    prompt_table = doc.add_table(rows = len(nameusedlist)+1 , cols = 2)
    prompt_table.rows[0].cells[0].text= "Prompts used"
    prompt_table.rows[0].cells[1].text= "Image generated"
    prompt_table.style = "Table Grid"
    for idx,eachimg in enumerate(nameusedlist):
        print(eachimg)
        prompt_table.cell(idx+1,0).text  = promptusedlist[idx]
        cell = prompt_table.cell(idx+1,1)
        img = cell.add_paragraph().add_run()
        img.add_picture(f"{report_path}/img/{eachimg}", width=Inches(2.5))
    # add an image to the first column of the table
    
    #code below for sustainability analysis
    doc.add_heading("Sustainability Recommendation Ranking" , level=2)
    answer = assess_sustainability.getchatgptranking(report_path)
    doc.add_paragraph(answer)
    
    #end code on sustainability analysis
    doc.add_heading("Analysis Conclusion" , level = 2)
    conclusion_string = f"""
    Thanks for using the product to generate a report for the {product}. We hope that the {min(20,len(reccotable))} recommendations 
    you have gotten was beneficial for your product developement journey! In general the product has received {globalsentiment}
    ratings from the youtube and amazon community which are our main source of feedback.
    """
    # add a paragraph of text to the document
    doc.add_paragraph(conclusion_string)
    # save the document
    doc.save(f'{report_path}/{product} recommendation report.docx')
